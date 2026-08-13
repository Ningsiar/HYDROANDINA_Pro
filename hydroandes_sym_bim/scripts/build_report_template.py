# -*- coding: utf-8 -*-
"""
scripts/build_report_template.py

Genera hydroandes_sym_bim/resources/plantilla_reporte.docx: la plantilla
Jinja2 (docxtpl) que usa core/report_generator_docxtpl.py para el reporte
Word "con plantilla y selección de secciones" de la Pestaña 21.

Este script NO se ejecuta en tiempo de uso del plugin -- es la fuente
versionada de la plantilla (un .docx es un binario, así que este .py es
lo que de verdad se revisa/edita cuando hay que cambiar el formato del
reporte). Se corre UNA VEZ, a mano, cada vez que se necesite regenerar la
plantilla:

    python-qgis.bat hydroandes_sym_bim/scripts/build_report_template.py

La estructura de secciones (portada, "1. DEM y Delimitación" ... "7.
Hidráulica y Drenaje") sigue el mismo orden y numeración que ya usaba
core/report_generator.py, para que ambos mecanismos de exportación (el
directo con python-docx y este con plantilla) resulten familiares al
mismo lector. Cada sección queda envuelta en un bloque condicional Jinja2
({% if incluir_x %} ... {% endif %}, cada tag en su propio párrafo para
que docxtpl los retire limpiamente) y con un marcador {{ subdoc_x }}
donde se inserta el contenido dinámico (tablas/gráficos) armado en
Python con la misma API de python-docx.
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

RUTA_SALIDA = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                            "resources", "plantilla_reporte.docx")

SECCIONES = [
    ("dem", "1. DEM y Delimitación de la Cuenca",
     "La cuenca se delimitó a partir de un Modelo de Elevación Digital (MDE), procesado mediante "
     "relleno de sumideros y direcciones/acumulación de flujo D8 (algoritmos GRASS), a partir de un "
     "punto de salida ajustado automáticamente al cauce más cercano."),
    ("morfometria", "2. Morfometría y Red de Drenaje",
     "Parámetros morfométricos e hidrológicos calculados a partir de la cuenca delimitada y el MDE "
     "recortado a su extensión, organizados en los grupos estándar de análisis morfométrico (básicos, "
     "forma, cauce principal, pendiente/curva hipsométrica, red de drenaje, y relieve/riesgo de flujo "
     "de detritos)."),
    ("cn", "3. Número de Curva SCS (CN)",
     "Número de curva bajo las tres condiciones antecedentes de humedad (AMC I, II y III), y la "
     "retención potencial máxima S y la abstracción inicial Ia asociadas al CN_II, base de las "
     "pérdidas por infiltración del método SCS-CN."),
    ("tc", "4. Tiempo de Concentración y Lag Time",
     "Tiempo de concentración estimado por métodos empíricos independientes, para comparar y adoptar "
     "el más representativo de las condiciones de la cuenca en estudio, junto con el coeficiente de "
     "escorrentía y la rugosidad de Manning cuando se calcularon con el catálogo de métodos del "
     "plugin."),
    ("frecuencia", "5. Precipitación Máxima en 24 Horas — Análisis de Frecuencia",
     "Ajuste de distribuciones de probabilidad a la serie de máximos anuales de P24h, con las pruebas "
     "de bondad de ajuste correspondientes, y las precipitaciones de diseño resultantes para cada "
     "periodo de retorno."),
    ("caudales", "6. Estimación de Caudales de Crecida",
     "Caudal pico e hidrograma de diseño obtenidos a partir de la tormenta de diseño y el modelo de "
     "transformación lluvia-escorrentía seleccionado (SCS, Snyder o Clark)."),
    ("hidraulica", "7. Hidráulica y Drenaje",
     "Dimensionamiento y verificación de las estructuras de drenaje calculadas en la sesión de "
     "trabajo (canales, alcantarillas, enrocado, sumideros, y verificaciones de borde libre de "
     "pontones/puentes/defensas ribereñas), con el cuadro comparativo final y la recomendación "
     "automática entre estructuras del mismo tipo."),
]


def _agregar_portada(doc: Document):
    titulo = doc.add_heading("HydroAndes SYM BIM — Reporte Técnico de Análisis Hidrológico e Hidráulico",
                              level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("{{ nombre_cuenca }}")
    r.bold = True
    r.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Generado: {{ fecha_generacion }}")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("Plugin HydroAndes SYM BIM — motor hidrológico/hidráulico para cuencas andinas del Perú")

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run("{{ nombre_empresa }}").italic = True

    doc.add_page_break()


def _agregar_seccion(doc: Document, clave: str, titulo: str, marco_teorico: str):
    doc.add_paragraph(f"{{% if incluir_{clave} %}}")
    doc.add_heading(titulo, level=1)
    doc.add_paragraph(marco_teorico)
    doc.add_paragraph(f"{{{{ subdoc_{clave} }}}}")
    doc.add_paragraph("{% endif %}")


def construir_plantilla():
    doc = Document()
    _agregar_portada(doc)
    for clave, titulo, marco_teorico in SECCIONES:
        _agregar_seccion(doc, clave, titulo, marco_teorico)

    doc.add_page_break()
    doc.add_heading("Nota", level=1)
    doc.add_paragraph(
        "Este reporte fue generado automáticamente por el plugin HydroAndes SYM BIM a partir de los "
        "valores calculados en la sesión de trabajo. Los métodos, coeficientes por defecto y "
        "supuestos empleados están documentados en el propio plugin; se recomienda verificar los "
        "resultados contra las fuentes normativas locales vigentes antes de un diseño definitivo."
    )

    os.makedirs(os.path.dirname(RUTA_SALIDA), exist_ok=True)
    doc.save(RUTA_SALIDA)
    return RUTA_SALIDA


if __name__ == "__main__":
    ruta = construir_plantilla()
    print(f"Plantilla generada en: {ruta}")
