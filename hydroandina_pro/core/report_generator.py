# -*- coding: utf-8 -*-
"""
core/report_generator.py

Genera un reporte técnico completo en Word (.docx) con los resultados de
las 7 pestañas del plugin (delimitación, morfometría, número de curva,
tiempo de concentración, precipitación, caudales, e hidráulica y
drenaje), con texto descriptivo, tablas y los gráficos ya generados en
la interfaz (se reciben como rutas a imágenes PNG ya exportadas por el
propio diálogo, para no volver a graficar nada aquí).

REQUIERE python-docx (no forma parte de qgis.core). Si no está
instalado en el intérprete de Python de QGIS:
    OSGeo4W Shell (Windows) o terminal con el Python de QGIS (Linux/Mac):
        pip install python-docx
Se importa de forma perezosa para no romper el resto del plugin si falta
la librería (mismo patrón que openpyxl en core/exporters.py).
"""
import os
from datetime import datetime
from typing import Optional


class ReportGeneratorError(Exception):
    pass


def _requerir_docx():
    try:
        import docx  # noqa: F401
        return docx
    except ImportError as e:
        raise ReportGeneratorError(
            "python-docx no está instalado en el intérprete de Python de QGIS. Ejecute "
            "'pip install python-docx' desde el OSGeo4W Shell (Windows) o el terminal con el "
            "Python de QGIS (Linux/Mac), y vuelva a intentar."
        ) from e


def _agregar_tabla_desde_filas(doc, encabezados, filas, estilo="Light Grid Accent 1"):
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    try:
        tabla.style = estilo
    except KeyError:
        pass  # si la plantilla de Word del usuario no tiene ese estilo, se deja el default
    hdr_cells = tabla.rows[0].cells
    for i, h in enumerate(encabezados):
        hdr_cells[i].text = str(h)
    for fila in filas:
        row_cells = tabla.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = "" if valor is None else str(valor)
    return tabla


def _agregar_imagen_si_existe(doc, ruta_imagen: Optional[str], ancho_cm=15.5):
    if ruta_imagen and os.path.exists(ruta_imagen):
        from docx.shared import Cm
        doc.add_picture(ruta_imagen, width=Cm(ancho_cm))


def generar_reporte_word(ruta_docx: str, contexto: dict) -> str:
    """
    contexto: dict con (todas las claves son opcionales; se omite la
    sección correspondiente si su clave no está presente o está vacía):
        nombre_cuenca: str
        break_point_xy: tuple
        morfometria_resultados: dict (g1,g2,g3,g4,g5,g6,lc_km)
        cn_resultados: dict (CN_I/II/III, S_mm, Ia_mm)
        desglose_cn: list de dicts (LULC x HSG), opcional
        tc_resultados: dict {nombre_metodo: {...}}
        resultados_frecuencia: dict {clave: {...}}
        mejor_ajuste_clave: str
        p24_disenio: dict {tr: p24}
        hidrograma_resultado: dict
        resultados_hidraulica_drenaje: dict {nombre_estructura: {...}}
        rutas_imagenes: dict con claves opcionales:
            'hipsometrica', 'frecuencia', 'comparacion_tr', 'comparacion_tr_cartesiano',
            'hidrograma', 'seccion_canal'
    """
    docx = _requerir_docx()
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- Portada ---
    titulo = doc.add_heading("HydroAndes Pro — Reporte Técnico de Análisis Hidrológico e Hidráulico", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Cuenca: {contexto.get('nombre_cuenca', '(sin nombre)')}\n").bold = True
    p.add_run(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n")
    p.add_run("Plugin HydroAndes Pro — motor hidrológico/hidráulico para cuencas andinas del Perú")
    doc.add_page_break()

    # ------------------------------------------------------------------
    # Pestaña 1: DEM y Delimitación
    # ------------------------------------------------------------------
    doc.add_heading("1. DEM y Delimitación de la Cuenca", level=1)
    doc.add_paragraph(
        "La cuenca se delimitó a partir de un Modelo de Elevación Digital (MDE), procesado "
        "mediante relleno de sumideros y direcciones/acumulación de flujo D8 (algoritmos GRASS), "
        "a partir de un punto de salida ajustado automáticamente al cauce más cercano."
    )
    bp = contexto.get("break_point_xy")
    if bp:
        doc.add_paragraph(f"Punto de salida (coordenadas del proyecto): X = {bp[0]:.2f}, Y = {bp[1]:.2f}")
    doc.add_paragraph(f"Cuenca activa: {contexto.get('nombre_cuenca', '(sin nombre)')}")

    # ------------------------------------------------------------------
    # Pestaña 2: Morfometría y drenaje
    # ------------------------------------------------------------------
    morfo = contexto.get("morfometria_resultados")
    if morfo:
        doc.add_heading("2. Morfometría y Red de Drenaje", level=1)
        doc.add_paragraph(
            "Parámetros morfométricos e hidrológicos calculados a partir de la cuenca delimitada "
            "y el MDE recortado a su extensión, organizados en los 6 grupos estándar de análisis "
            "morfométrico (básicos, forma, cauce principal, pendiente/curva hipsométrica, red de "
            "drenaje, y relieve/riesgo de flujo de detritos)."
        )
        try:
            from .plugin_dialog import NOMBRES_PARAMETROS_MORFOMETRIA, UNIDADES_PARAMETROS_MORFOMETRIA
        except Exception:
            NOMBRES_PARAMETROS_MORFOMETRIA, UNIDADES_PARAMETROS_MORFOMETRIA = {}, {}
        filas = []
        for grupo_key in ("g1", "g2", "g3", "g5", "g6"):
            grupo = morfo.get(grupo_key, {})
            for k, val in grupo.items():
                if k in ("interpretacion", "curva_hipsometrica", "mensaje_alerta") or isinstance(val, (list, dict)):
                    continue
                nombre = NOMBRES_PARAMETROS_MORFOMETRIA.get(k, k)
                unidad = UNIDADES_PARAMETROS_MORFOMETRIA.get(k, "")
                valor_str = f"{val:.4f}" if isinstance(val, float) else str(val)
                filas.append((nombre, k, valor_str, unidad))
        if filas:
            _agregar_tabla_desde_filas(doc, ["Parámetro", "Símbolo", "Valor", "Unidad"], filas)

        g4 = morfo.get("g4")
        if g4 and contexto.get("rutas_imagenes", {}).get("hipsometrica"):
            doc.add_heading("Curva hipsométrica", level=2)
            doc.add_paragraph(f"Pendiente media de la cuenca: {g4.get('S_cuenca_pct', '?')} %")
            _agregar_imagen_si_existe(doc, contexto["rutas_imagenes"]["hipsometrica"])

        cav = contexto.get("resultado_cav")
        if cav:
            doc.add_heading("Curva Cota-Área-Volumen (C-A-V)", level=2)
            doc.add_paragraph(
                f"Cota mínima: {cav.get('z_min')} m s.n.m.   Cota máxima: {cav.get('z_max')} m s.n.m.   "
                f"Área total: {cav.get('area_total_m2', 0) / 1e6:.4f} km²   "
                f"Volumen total embalsable: {cav.get('volumen_total_m3', 0) / 1e6:.4f} hm³"
            )
            _agregar_imagen_si_existe(doc, contexto.get("rutas_imagenes", {}).get("cav"))

    # ------------------------------------------------------------------
    # Pestaña 3: Número de Curva SCS
    # ------------------------------------------------------------------
    cn = contexto.get("cn_resultados")
    if cn:
        doc.add_heading("3. Número de Curva SCS (CN)", level=1)
        doc.add_paragraph(
            "Número de curva bajo las tres condiciones antecedentes de humedad (AMC I, II y III), "
            "y la retención potencial máxima S y la abstracción inicial Ia asociadas al CN_II."
        )
        _agregar_tabla_desde_filas(
            doc, ["CN_I", "CN_II", "CN_III", "S (mm)", "Ia (mm)"],
            [(cn.get("CN_I"), cn.get("CN_II"), cn.get("CN_III"), cn.get("S_mm"), cn.get("Ia_mm"))]
        )
        desglose_cn = contexto.get("desglose_cn")
        if desglose_cn:
            doc.add_heading("Desglose por uso de suelo / grupo hidrológico", level=2)
            filas = [(d.get("lulc_nombre", d.get("uso", "")), d.get("hsg", ""), d.get("area_km2"), d.get("cn"))
                     for d in desglose_cn]
            _agregar_tabla_desde_filas(doc, ["Uso de suelo", "Grupo hidrológico", "Área (km²)", "CN"], filas)

    # ------------------------------------------------------------------
    # Pestaña 4: Tiempo de concentración y Lag Time
    # ------------------------------------------------------------------
    tc = contexto.get("tc_resultados")
    if tc:
        doc.add_heading("4. Tiempo de Concentración y Lag Time", level=1)
        doc.add_paragraph(
            "Tiempo de concentración estimado por 14 métodos empíricos distintos, para comparar y "
            "adoptar el más representativo de las condiciones de la cuenca en estudio."
        )
        filas = [(nombre, datos.get("autor_anio", ""),
                   datos.get("Tc_horas") if not datos.get("error") else f"ERROR: {datos['error']}",
                   datos.get("Tc_min", ""), datos.get("tlag_min", ""))
                  for nombre, datos in tc.items()]
        _agregar_tabla_desde_filas(doc, ["Método", "Autor/Año", "Tc (h)", "Tc (min)", "tlag SCS (min)"], filas)

    # ------------------------------------------------------------------
    # Pestaña 5: Precipitación Máxima 24h
    # ------------------------------------------------------------------
    freq = contexto.get("resultados_frecuencia")
    if freq:
        doc.add_heading("5. Precipitación Máxima en 24 Horas — Análisis de Frecuencia", level=1)
        doc.add_paragraph(
            "Ajuste de distribuciones de probabilidad a la serie de máximos anuales de P24h, con "
            "la prueba de bondad de ajuste de Kolmogorov-Smirnov, y las precipitaciones de diseño "
            "resultantes para cada periodo de retorno."
        )
        mejor = contexto.get("mejor_ajuste_clave")
        filas = []
        for clave, r in freq.items():
            if r.get("error"):
                continue
            marca = "  ★ mejor ajuste" if clave == mejor else ""
            params_str = ", ".join(f"{k}={v:.4f}" for k, v in r["parametros"].items())
            filas.append((r["nombre"] + marca, params_str, r["D_ks"], r["D_critico"], "Sí" if r["pasa_ks"] else "No"))
        if filas:
            _agregar_tabla_desde_filas(doc, ["Distribución", "Parámetros", "D (KS)", "D crítico", "¿Pasa KS?"], filas)

        p24 = contexto.get("p24_disenio")
        if p24:
            doc.add_heading("Precipitaciones de diseño (mejor ajuste)", level=2)
            _agregar_tabla_desde_filas(
                doc, [f"Tr={tr}a" for tr in p24.keys()], [[p24[tr] for tr in p24.keys()]]
            )

        rutas_img = contexto.get("rutas_imagenes", {})
        for clave_img, titulo_img in [("frecuencia", "Ajuste de distribuciones"),
                                        ("comparacion_tr", "Comparación Pmax24h vs Tr (escala log)"),
                                        ("comparacion_tr_cartesiano", "Comparación Pmax24h vs Tr (escala cartesiana)")]:
            if rutas_img.get(clave_img):
                doc.add_heading(titulo_img, level=2)
                _agregar_imagen_si_existe(doc, rutas_img[clave_img])

    # ------------------------------------------------------------------
    # Pestaña 6: Caudales SCS/Snyder/Clark
    # ------------------------------------------------------------------
    hidrograma = contexto.get("hidrograma_resultado")
    if hidrograma:
        doc.add_heading("6. Estimación de Caudales de Crecida", level=1)
        doc.add_paragraph(
            f"Caudal pico de diseño: Qp = {hidrograma.get('caudal_pico_m3s')} m³/s, "
            f"en el tiempo Tp = {hidrograma.get('tiempo_pico_h')} h."
        )
        if contexto.get("rutas_imagenes", {}).get("hidrograma"):
            _agregar_imagen_si_existe(doc, contexto["rutas_imagenes"]["hidrograma"])

    # ------------------------------------------------------------------
    # Pestaña 7: Hidráulica y Drenaje
    # ------------------------------------------------------------------
    hidraulica = contexto.get("resultados_hidraulica_drenaje")
    if hidraulica:
        doc.add_heading("7. Hidráulica y Drenaje", level=1)
        doc.add_paragraph(
            "Dimensionamiento y verificación de las estructuras de drenaje calculadas en esta "
            "sesión (canales, alcantarillas, enrocado, sumideros, y verificaciones de borde libre "
            "de pontones/puentes/defensas ribereñas)."
        )
        for nombre_estructura, datos in hidraulica.items():
            doc.add_heading(nombre_estructura, level=2)
            filas = [(k, v) for k, v in datos.items() if k != "tipo"]
            _agregar_tabla_desde_filas(doc, ["Parámetro/Resultado", "Valor"], filas)

    # --- Pie de página / advertencia general ---
    doc.add_page_break()
    doc.add_heading("Nota", level=1)
    doc.add_paragraph(
        "Este reporte fue generado automáticamente por el plugin HydroAndes Pro a partir de los "
        "valores calculados en la sesión de trabajo. Los métodos, coeficientes por defecto y "
        "supuestos empleados están documentados en el propio plugin; se recomienda verificar los "
        "resultados contra las fuentes normativas locales vigentes antes de un diseño definitivo."
    )

    os.makedirs(os.path.dirname(ruta_docx) or ".", exist_ok=True)
    doc.save(ruta_docx)
    return ruta_docx
